"""LocalFileCollector — reads local files (txt, md, pdf, docx) into Documents."""

from __future__ import annotations

from pathlib import Path

from loguru import logger

from src.config import Document, ResearchConfig

from .base import BaseCollector

# Поддерживаемые расширения
_SUPPORTED_EXT = {".txt", ".md", ".pdf", ".docx"}


class LocalFileCollector(BaseCollector):
    """Collect documents from local filesystem paths."""

    async def collect(
        self, config: ResearchConfig
    ) -> tuple[list[Document], list[str], int]:
        """
        config.source_urls содержит пути к файлам (не URL).
        Возвращает (documents, failed_paths, 0).
        blocked_count всегда 0 — стоп-листа для файлов нет.
        """
        paths = list(config.source_urls)
        docs: list[Document] = []
        failed: list[str] = []
        seen_hashes: set[str] = set()

        for path_str in paths:
            path = Path(path_str)
            try:
                doc = self._read_file(path)
                if doc is None:
                    failed.append(path_str)
                elif doc.content_hash in seen_hashes:
                    logger.debug("Дубликат по content_hash: {}", path_str)
                else:
                    seen_hashes.add(doc.content_hash)
                    docs.append(doc)
            except Exception as exc:
                logger.warning("Ошибка чтения {}: {}", path_str, exc)
                failed.append(path_str)

        logger.info(
            "LocalFileCollector: {} документов, {} ошибок для темы '{}'",
            len(docs), len(failed), config.topic,
        )
        return docs, failed, 0

    def _read_file(self, path: Path) -> Document | None:
        """Прочитать файл и вернуть Document. None если неподдерживаемый тип."""
        if not path.exists():
            logger.warning("Файл не найден: {}", path)
            return None

        ext = path.suffix.lower()

        if ext not in _SUPPORTED_EXT:
            logger.debug("Неподдерживаемый тип файла: {}", ext)
            return None

        if ext in {".txt", ".md"}:
            text = self._read_text(path)
        elif ext == ".pdf":
            text = self._read_pdf(path)
        elif ext == ".docx":
            text = self._read_docx(path)
        else:
            return None

        if not text or len(text.strip()) < 50:
            logger.debug("Слишком мало контента в {}", path)
            return None

        return Document(
            url=str(path),
            title=path.stem,
            content=text.strip(),
        )

    @staticmethod
    def _read_text(path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    @staticmethod
    def _read_pdf(path: Path) -> str:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber не установлен: pip install pdfplumber")

        pages: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n\n".join(pages)

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx не установлен: pip install python-docx")

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
